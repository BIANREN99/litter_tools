import sys
import csv
import random
import re
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                            QHBoxLayout, QLabel, QLineEdit, QPushButton, 
                            QTextEdit, QFileDialog, QMessageBox, QTabWidget,
                            QProgressBar, QGroupBox)
from PyQt5.QtCore import Qt
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PyQt5.QtWidgets import QRadioButton, QButtonGroup, QSpinBox, QLabel
class WordLearningTool(QMainWindow):
    def __init__(self):
        super().__init__()
        self.words = []
        self.learned_words = set()
        self.current_round_words = []
        self.current_word = None
        self.answer_shown = False
        # 修改为两个独立的已导出单词集合
        self.exported_words_eng_to_chi = set()  # 英译中模式的已导出单词
        self.exported_words_chi_to_eng = set()  # 中译英模式的已导出单词
        self.initUI()

    def initUI(self):
        self.setWindowTitle('英语单词学习工具')
        self.setGeometry(100, 100, 800, 600)
        
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        
        tabs = QTabWidget()
        main_layout.addWidget(tabs)
        
        learning_tab = QWidget()
        tabs.addTab(learning_tab, "单词学习")
        self.initLearningTab(learning_tab)
        
        export_tab = QWidget()
        tabs.addTab(export_tab, "导出Word")
        self.initExportTab(export_tab)
        
        self.statusBar().showMessage('就绪')
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        main_layout.addWidget(self.progress_bar)
        
    def initLearningTab(self, parent):
        layout = QVBoxLayout(parent)
        
        # 文件操作组
        file_group = QGroupBox("文件操作")
        file_layout = QHBoxLayout(file_group)
        
        self.file_path_edit = QLineEdit()
        self.file_path_edit.setPlaceholderText("请选择单词文件...")
        file_layout.addWidget(self.file_path_edit)
        
        browse_btn = QPushButton("浏览...")
        browse_btn.clicked.connect(self.browse_file)
        file_layout.addWidget(browse_btn)
        
        load_btn = QPushButton("加载单词")
        load_btn.clicked.connect(self.load_words)
        file_layout.addWidget(load_btn)
        
        layout.addWidget(file_group)
        
        # 学习模式选择
        mode_group = QGroupBox("学习模式")
        mode_layout = QHBoxLayout(mode_group)
        
        self.english_to_chinese_btn = QPushButton("英译中")
        self.english_to_chinese_btn.clicked.connect(self.start_english_to_chinese)
        mode_layout.addWidget(self.english_to_chinese_btn)
        
        self.chinese_to_english_btn = QPushButton("中译英")
        self.chinese_to_english_btn.clicked.connect(self.start_chinese_to_english)
        mode_layout.addWidget(self.chinese_to_english_btn)
        
        layout.addWidget(mode_group)
        
        # 学习区域
        study_group = QGroupBox("学习区域")
        study_layout = QVBoxLayout(study_group)
        
        self.word_label = QLabel("点击上方按钮开始学习")
        self.word_label.setAlignment(Qt.AlignCenter)
        self.word_label.setStyleSheet("font-size: 24px; margin: 20px;")
        study_layout.addWidget(self.word_label)
        
        self.input_edit = QLineEdit()
        self.input_edit.setPlaceholderText("在此输入答案...")
        self.input_edit.returnPressed.connect(self.check_answer)
        study_layout.addWidget(self.input_edit)
        
        button_layout = QHBoxLayout()
        
        self.check_btn = QPushButton("检查答案")
        self.check_btn.clicked.connect(self.check_answer)
        button_layout.addWidget(self.check_btn)
        
        # 添加显示答案按钮
        self.show_answer_btn = QPushButton("显示答案")
        self.show_answer_btn.clicked.connect(self.show_answer)
        button_layout.addWidget(self.show_answer_btn)
        
        self.next_btn = QPushButton("下一个")
        self.next_btn.clicked.connect(self.next_word)
        button_layout.addWidget(self.next_btn)
        
        study_layout.addLayout(button_layout)
        
        self.result_label = QLabel("")
        self.result_label.setAlignment(Qt.AlignCenter)
        self.result_label.setStyleSheet("font-size: 18px; margin: 10px;")
        study_layout.addWidget(self.result_label)
        
        layout.addWidget(study_group)
        
        # 统计信息
        stats_group = QGroupBox("学习统计")
        stats_layout = QHBoxLayout(stats_group)
        
        self.total_label = QLabel("总单词数: 0")
        stats_layout.addWidget(self.total_label)
        
        self.learned_label = QLabel("已学习: 0")
        stats_layout.addWidget(self.learned_label)
        
        self.remaining_label = QLabel("剩余: 0")
        stats_layout.addWidget(self.remaining_label)
        
        layout.addWidget(stats_group)
        
    def initExportTab(self, parent):
        layout = QVBoxLayout(parent)
        
        export_group = QGroupBox("导出设置")
        export_layout = QVBoxLayout(export_group)
        
        # 导出模式选择
        mode_layout = QHBoxLayout()
        mode_label = QLabel("导出模式：")
        mode_layout.addWidget(mode_label)
        
        self.export_mode_group = QButtonGroup(self)
        
        self.eng_to_chi_radio = QRadioButton("英译中")
        self.eng_to_chi_radio.setChecked(True)  # 默认选择英译中
        self.chi_to_eng_radio = QRadioButton("中译英")
        
        self.export_mode_group.addButton(self.eng_to_chi_radio)
        self.export_mode_group.addButton(self.chi_to_eng_radio)
        mode_layout.addWidget(self.eng_to_chi_radio)
        mode_layout.addWidget(self.chi_to_eng_radio)
        mode_layout.addStretch()
        export_layout.addLayout(mode_layout)
        
        # 单词数量选择
        count_layout = QHBoxLayout()
        count_label = QLabel("每套装单词数量：")
        count_layout.addWidget(count_label)
        
        self.word_count_spin = QSpinBox()
        self.word_count_spin.setMinimum(1)
        self.word_count_spin.setMaximum(1000)
        self.word_count_spin.setValue(50)  # 默认50个单词
        self.word_count_spin.setSuffix(" 个")
        count_layout.addWidget(self.word_count_spin)
        count_layout.addStretch()
        export_layout.addLayout(count_layout)
        
        # 套数选择
        sets_layout = QHBoxLayout()
        sets_label = QLabel("导出套数：")
        sets_layout.addWidget(sets_label)
        
        self.sets_spin = QSpinBox()
        self.sets_spin.setMinimum(1)
        self.sets_spin.setMaximum(20)
        self.sets_spin.setValue(1)  # 默认1套
        self.sets_spin.setSuffix(" 套")
        sets_layout.addWidget(self.sets_spin)
        sets_layout.addStretch()
        export_layout.addLayout(sets_layout)
        
        # 导出路径设置
        self.export_path_edit = QLineEdit()
        self.export_path_edit.setPlaceholderText("选择导出位置...")
        export_layout.addWidget(self.export_path_edit)
        
        browse_export_btn = QPushButton("选择导出位置...")
        browse_export_btn.clicked.connect(self.browse_export_path)
        export_layout.addWidget(browse_export_btn)
        
        export_btn = QPushButton("导出Word文档")
        export_btn.clicked.connect(self.export_to_word)
        export_layout.addWidget(export_btn)
        
        layout.addWidget(export_group)
        
        info_text = QTextEdit()
        info_text.setReadOnly(True)
        info_text.setPlainText("""
导出说明：
1. 选择导出模式：英译中（显示英文单词，留空中文）或中译英（显示中文意思，留空英文）
2. 设置每套装的单词数量和要导出的套数
3. 点击"选择导出位置"按钮选择保存位置
4. 点击"导出Word文档"按钮开始导出
5. 系统将随机选择指定数量的单词进行导出
6. 每个模式有独立的已使用单词记录，两个模式的单词使用互不影响
7. 当模式下的所有单词都导出过一轮后，会重置导出状态并提示用户
8. 多套导出时，会在文件名后自动添加序号（如：words_dictation_1.docx）
9. 建议在导出前确保已加载单词文件
        """)
        layout.addWidget(info_text)
        
    def browse_file(self):
        file_name, _ = QFileDialog.getOpenFileName(
            self,
            "选择单词文件",
            "",
            "文本文件 (*.txt);;CSV文件 (*.csv);;Markdown文件 (*.md)"
        )
        
        if file_name:
            self.file_path_edit.setText(file_name)
            self.file_path_edit.repaint()
    
    def browse_export_path(self):
        file_name, _ = QFileDialog.getSaveFileName(
            self,
            "选择导出位置",
            "words_dictation.docx",
            "Word文档 (*.docx)"
        )
        if file_name:
            self.export_path_edit.setText(file_name)
            
    def load_words(self, file_path: str = None) -> bool:
        if file_path is None or file_path is False:  # ★ 防止接收到 False
            file_path = self.file_path_edit.text().strip()
        
        if not file_path:
            QMessageBox.warning(self, "警告", "请先选择单词文件！")
            return False
        
        if not file_path.lower().endswith(('.txt', '.csv', '.md')):
            QMessageBox.warning(self, "格式错误", "仅支持.txt、.csv和.md格式的文件")
            return False
        
        self.words = []
        
        try:
            if file_path.lower().endswith('.csv'):
                # 现有CSV处理逻辑保持不变
                with open(file_path, 'r', encoding='utf-8') as f:
                    csv_reader = csv.reader(f)
                    headers = next(csv_reader, None)  # 跳过标题行
                    for row in csv_reader:
                        if len(row) >= 3:  # 确保至少有三列数据
                            word = row[0].strip()  # 英文单词
                            meaning = row[2].strip()  # 中文意思
                            
                            # 验证word是否为英文，meaning是否包含中文
                            is_english_word = bool(re.match(r'^[a-zA-Z\s\-\']+$', word))
                            has_chinese_meaning = any('\u4e00' <= c <= '\u9fff' for c in meaning)
                            
                            if is_english_word and has_chinese_meaning:
                                self.words.append({'english': word, 'chinese': meaning})
                            elif has_chinese_meaning:  # 如果word不是明显的英文但meaning包含中文，视为中文在前
                                self.words.append({'chinese': word, 'english': meaning})
                            else:  # 保守处理，默认英文在前
                                self.words.append({'english': word, 'chinese': meaning})
            elif file_path.lower().endswith('.md'):
                # 添加Markdown表格解析逻辑
                with open(file_path, 'r', encoding='utf-8') as f:
                    lines = f.readlines()
                    
                    # 跳过表头和分隔行
                    table_started = False
                    for line in lines:
                        line = line.strip()
                        # 跳过空行
                        if not line:
                            continue
                        
                        # 检查是否是表格分隔行（包含|----|----|----|格式）
                        if '|' in line and all(cell.strip() == '' or cell.strip() == '-'*len(cell.strip()) for cell in line.split('|')):
                            table_started = True
                            continue
                        
                        # 只处理表格数据行
                        if table_started and line.startswith('|'):
                            # 分割表格单元格
                            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                            
                            # 确保至少有三列数据
                            if len(cells) >= 3:
                                word = cells[0]  # 英文单词
                                meaning = cells[2]  # 中文意思
                                
                                # 验证word是否为英文，meaning是否包含中文
                                is_english_word = bool(re.match(r'^[a-zA-Z\s\-\']+$', word))
                                has_chinese_meaning = any('\u4e00' <= c <= '\u9fff' for c in meaning)
                                
                                if is_english_word and has_chinese_meaning:
                                    self.words.append({'english': word, 'chinese': meaning})
                                elif has_chinese_meaning:  # 如果word不是明显的英文但meaning包含中文，视为中文在前
                                    self.words.append({'chinese': word, 'english': meaning})
                                else:  # 保守处理，默认英文在前
                                    self.words.append({'english': word, 'chinese': meaning})
            else:  # .txt文件
                with open(file_path, 'r', encoding='utf-8') as f:
                    for line in f:
                        parsed = self.parse_line(line.strip())
                        if parsed:
                            self.words.append(parsed)
                
            if not self.words:
                QMessageBox.warning(self, "解析警告", "未能成功解析任何单词，请检查文件格式")
                return False
                
            self.reset_round()
            self.update_stats()
            QMessageBox.information(self, "成功", f"成功加载 {len(self.words)} 个单词！")
            return True
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"加载文件时出错：{str(e)}")
            return False
            
    def reset_round(self):
        self.current_round_words = [
            w for w in self.words if w['english'] not in self.learned_words
        ]
        random.shuffle(self.current_round_words)
        
    def update_stats(self):
        self.total_label.setText(f"总单词数: {len(self.words)}")
        self.learned_label.setText(f"已学习: {len(self.learned_words)}")
        self.remaining_label.setText(f"剩余: {len(self.current_round_words)}")
        
    def start_english_to_chinese(self):
        self.current_mode = 'english_to_chinese'
        self.next_word()
        
    def start_chinese_to_english(self):
        self.current_mode = 'chinese_to_english'
        self.next_word()
        
    def next_word(self):
        if not self.current_round_words:
            self.reset_round()
            
        if self.current_round_words:
            self.current_word = self.current_round_words.pop(0)
            self.input_edit.clear()
            self.result_label.clear()
            self.answer_shown = False
            self.input_edit.setEnabled(True)
            
            if self.current_mode == 'english_to_chinese':
                self.word_label.setText(f"英文单词：\n{self.current_word['english']}")
                self.input_edit.setPlaceholderText("请输入中文意思...")
            else:
                self.word_label.setText(f"中文意思：\n{self.current_word['chinese']}")
                self.input_edit.setPlaceholderText("请输入英文单词...")
                
            self.update_stats()
        else:
            self.word_label.setText("本轮学习已完成！")
            self.input_edit.setEnabled(False)
            
    def check_answer(self):
        if not self.current_word:
            return
            
        user_input = self.input_edit.text().strip()
        if not user_input:
            return
            
        if self.current_mode == 'english_to_chinese':
            # 英译中模式：检查用户输入是否包含在正确答案中
            # 转换为小写进行比较，忽略大小写差异
            user_input_lower = user_input.lower()
            correct_answer_lower = self.current_word['chinese'].lower()
            correct = user_input_lower in correct_answer_lower
        else:
            correct = user_input.lower() == self.current_word['english'].lower()
            
        if correct:
            self.result_label.setText("✓ 回答正确！答案是：{answer}")
            self.result_label.setStyleSheet("color: green; font-size: 18px;")
            self.learned_words.add(self.current_word['english'])
        else:
            answer = self.current_word['chinese'] if self.current_mode == 'english_to_chinese' else self.current_word['english']
            self.result_label.setText(f"✗ 回答错误！正确答案是：{answer}")
            self.result_label.setStyleSheet("color: red; font-size: 18px;")
            
        self.update_stats()
    # 添加显示答案方法
    def show_answer(self):
        if not self.current_word:
            return
            
        if self.answer_shown:
            # 如果答案已经显示，则清除显示
            self.result_label.setText("")
            self.result_label.setStyleSheet("")
            self.answer_shown = False
        else:
            # 如果答案未显示，则显示答案
            answer = self.current_word['chinese'] if self.current_mode == 'english_to_chinese' else self.current_word['english']
            self.result_label.setText(f"答案：{answer}")
            self.result_label.setStyleSheet("color: blue; font-size: 18px;")
            self.answer_shown = True
    def parse_line(self, line):
        """解析文件中的每一行，提取英文单词和中文意思
        支持多种格式，包括：单词|词性|中文意思
        """
        # 首先检查是否是竖线分隔的格式：单词|词性|中文意思
        if '|' in line:
            parts = line.split('|')
            if len(parts) >= 3:
                # 提取英语单词和中文意思，忽略词性部分
                english_word = parts[0].strip()
                # 中文意思可能包含多个部分，合并后面的所有部分
                chinese_meanings = '|'.join(parts[2:]).strip()
                
                # 验证这确实是英文单词和中文意思
                if not any('\u4e00' <= c <= '\u9fff' for c in english_word) and chinese_meanings:
                    return {'english': english_word, 'chinese': chinese_meanings}
        
        # 原有的分隔符支持保持不变
        separators = [',', '：', ':', '\t', '  ']
        
        for sep in separators:
            if sep in line:
                parts = line.split(sep, 1)
                if len(parts) == 2:
                    part1, part2 = parts[0].strip(), parts[1].strip()
                    if any('\u4e00' <= c <= '\u9fff' for c in part1):
                        return {'chinese': part1, 'english': part2}
                    else:
                        return {'english': part1, 'chinese': part2}
        
        # 空格分隔的情况保持不变
        words = line.split()
        if len(words) >= 2:
            english = []
            chinese = []
            for w in words:
                if any('\u4e00' <= c <= '\u9fff' for c in w):
                    chinese.append(w)
                else:
                    english.append(w)
            if english and chinese:
                return {'english': ' '.join(english), 'chinese': ' '.join(chinese)}
        
        return None

    def export_to_word(self):
        export_path = self.export_path_edit.text()
        if not export_path:
            QMessageBox.warning(self, "警告", "请先选择导出位置！")
            return
            
        if not self.words:
            QMessageBox.warning(self, "警告", "请先加载单词文件！")
            return
        
        try:
            # 获取导出设置
            is_english_to_chinese = self.eng_to_chi_radio.isChecked()
            word_count = self.word_count_spin.value()
            sets_count = self.sets_spin.value()
            
            # 检查单词数量是否合理
            total_words = len(self.words)
            if word_count > total_words:
                QMessageBox.warning(self, "警告", f"单词总数不足！当前只有{total_words}个单词，无法导出{word_count}个单词。")
                return
            
            # 检查套数是否合理
            max_possible_sets = total_words // word_count
            if sets_count > max_possible_sets:
                # 计算在不重复的情况下最多能导出多少套
                response = QMessageBox.question(self, "提示", 
                    f"在不重复单词的情况下，最多只能导出{max_possible_sets}套。\n"
                    f"是否继续导出{sets_count}套（部分单词将重复）？",
                    QMessageBox.Yes | QMessageBox.No)
                if response == QMessageBox.No:
                    return
            
            # 根据模式选择对应的已导出单词集合
            if is_english_to_chinese:
                exported_words = self.exported_words_eng_to_chi
            else:
                exported_words = self.exported_words_chi_to_eng
            
            # 保存原始导出路径
            import os
            base_name, ext = os.path.splitext(export_path)
            
            # 导出多套
            for set_num in range(1, sets_count + 1):
                # 为每套装构建导出路径
                if sets_count > 1:
                    current_export_path = f"{base_name}_{set_num}{ext}"
                else:
                    current_export_path = export_path
                
                # 获取可用单词（未导出过的或所有单词，如果已全部导出）
                available_words = [word for word in self.words if word['english'] not in exported_words]
                
                # 如果没有可用单词（都已导出过），重置已导出集合
                need_reset = False
                if not available_words:
                    exported_words.clear()
                    available_words = self.words.copy()
                    need_reset = True
                
                # 检查可用单词数量是否足够
                if len(available_words) < word_count:
                    # 如果可用单词不够，从所有单词中补充
                    remaining_needed = word_count - len(available_words)
                    all_words_except_used = [word for word in self.words if word not in available_words]
                    
                    # 如果还有单词可以补充
                    if all_words_except_used:
                        # 随机选择需要补充的单词
                        additional_words = random.sample(all_words_except_used, min(remaining_needed, len(all_words_except_used)))
                        available_words.extend(additional_words)
                    
                    # 如果还是不够，从所有单词中随机选择（允许重复）
                    if len(available_words) < word_count:
                        remaining_needed = word_count - len(available_words)
                        random_words = random.sample(self.words, remaining_needed)
                        available_words.extend(random_words)
                
                # 随机选择指定数量的单词
                selected_words = random.sample(available_words, word_count)
                
                # 更新已导出单词集合
                for word in selected_words:
                    exported_words.add(word['english'])
                
                # 创建Word文档
                doc = Document()
                
                # 设置标题
                if is_english_to_chinese:
                    title_text = f'英语单词默写（英译中）- 第{set_num}套'
                    if sets_count == 1:
                        title_text = '英语单词默写（英译中）'
                    title = doc.add_heading(title_text, 0)
                    doc.add_paragraph('请根据英文写出对应的中文意思：')
                else:
                    title_text = f'英语单词默写（中译英）- 第{set_num}套'
                    if sets_count == 1:
                        title_text = '英语单词默写（中译英）'
                    title = doc.add_heading(title_text, 0)
                    doc.add_paragraph('请根据中文写出对应的英文单词：')
                
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                doc.add_paragraph('')
                
                # 添加信息段落
                info_p = doc.add_paragraph()
                info_p.add_run(f'单词数量：{word_count}个').bold = True
                if need_reset:
                    info_p.add_run(' （已重置单词列表）')
                doc.add_paragraph('')
                
                # 创建表格
                table = doc.add_table(rows=0, cols=2)
                table.style = 'Table Grid'
                
                for word in selected_words:
                    row = table.add_row()
                    if is_english_to_chinese:
                        # 英译中模式：显示英文，留空中文
                        row.cells[0].text = word['english']
                        row.cells[1].text = '__________'
                    else:
                        # 中译英模式：显示中文，留空英文
                        row.cells[0].text = word['chinese']
                        row.cells[1].text = '__________'
                        
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            p.runs[0].font.size = Pt(12)
                
                # 添加答案页
                doc.add_page_break()
                doc.add_heading('答案', level=1)
                
                for word in selected_words:
                    p = doc.add_paragraph()
                    if is_english_to_chinese:
                        # 英译中答案：英文 + 中文
                        p.add_run(word['english'] + ': ').bold = True
                        p.add_run(word['chinese'])
                    else:
                        # 中译英答案：中文 + 英文
                        p.add_run(word['chinese'] + ': ').bold = True
                        p.add_run(word['english'])
                
                # 保存文档
                doc.save(current_export_path)
            
            # 更新对应模式的已导出单词集合
            if is_english_to_chinese:
                self.exported_words_eng_to_chi = exported_words
                remaining_count = len([word for word in self.words if word['english'] not in exported_words])
                mode_name = "英译中"
            else:
                self.exported_words_chi_to_eng = exported_words
                remaining_count = len([word for word in self.words if word['english'] not in exported_words])
                mode_name = "中译英"
            
            # 显示成功消息
            message = f"Word文档已成功导出！\n"
            message += f"本次共导出{sets_count}套，每套{word_count}个单词。\n"
            message += f"{mode_name}模式下还剩余{remaining_count}个单词可导出。"
            
            QMessageBox.information(self, "成功", message)
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出时出错：{str(e)}")

def main():
    app = QApplication(sys.argv)
    tool = WordLearningTool()
    tool.show()
    sys.exit(app.exec_())

if __name__ == '__main__':
    main()